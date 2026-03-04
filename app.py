import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import io

st.set_page_config(
    page_title="Consulting Research Agent",
    page_icon="📊",
    layout="wide"
)

# ---- HEADER ----
st.markdown("""
<style>
    .main-header {
        background-color: #1F3564;
        padding: 20px 30px;
        border-radius: 8px;
        margin-bottom: 24px;
    }
    .main-header h1 { color: white; margin: 0; font-size: 26px; }
    .main-header p { color: #C0C0C0; margin: 4px 0 0 0; font-size: 14px; }
    .section-header {
        background-color: #EAF2FB;
        padding: 8px 14px;
        border-left: 4px solid #1F3564;
        border-radius: 4px;
        margin: 16px 0 8px 0;
        font-weight: bold;
        color: #1F3564;
    }
    .stButton>button {
        background-color: #1F3564;
        color: white;
        font-size: 16px;
        padding: 12px 40px;
        border-radius: 6px;
        border: none;
        width: 100%;
    }
    .stButton>button:hover { background-color: #2E74B5; }
</style>
<div class="main-header">
    <h1>📊 Consulting Research Agent</h1>
    <p>Generate structured research and analysis for client meetings</p>
</div>
""", unsafe_allow_html=True)

# ---- SIDEBAR: API KEY ----
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    api_key = st.text_input("Gemini API Key", type="password",
                             help="Enter your Gemini API key. Get one free at aistudio.google.com")
    st.markdown("---")
    st.markdown("**How to use:**")
    st.markdown("1. Enter your Gemini API key")
    st.markdown("2. Fill in the research inputs")
    st.markdown("3. Select sections needed")
    st.markdown("4. Click Generate")
    st.markdown("5. Download your document")
    st.markdown("---")
    st.markdown("*Research is based on Gemini's training data. For live data, web search upgrade is coming.*")

# ---- MAIN FORM ----
st.markdown('<div class="section-header">🏢 Company & Meeting Details</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    company = st.text_input("Company Name *", placeholder="e.g. Maruti Suzuki")
    industry = st.text_input("Industry *", placeholder="e.g. Indian Passenger Vehicles")
    meeting_topic = st.text_input("Meeting Topic *",
                                   placeholder="e.g. Improving customer experience")

with col2:
    meeting_objective = st.selectbox(
        "Meeting Objective *",
        ["Growth", "Cost Reduction", "Transformation", "New Product Development",
         "Customer Experience", "Digital Strategy", "Market Expansion", "Other"]
    )
    competitors_raw = st.text_input(
        "Competitors (comma separated) *",
        placeholder="e.g. Hyundai India, Tata Motors, Mahindra"
    )

st.markdown('<div class="section-header">🌍 Geography & Context</div>', unsafe_allow_html=True)

col3, col4 = st.columns(2)
with col3:
    geography_option = st.selectbox(
        "Geography Focus",
        ["India Only", "India + Global", "India + Specific Countries"]
    )
    if geography_option == "India + Specific Countries":
        specific_countries = st.text_input(
            "Specify Countries",
            placeholder="e.g. USA, Germany, China, Japan"
        )
    else:
        specific_countries = ""

with col4:
    output_format = st.selectbox(
        "Output Format",
        ["Word Document (.docx)", "PowerPoint (.pptx)", "Both"]
    )

context_notes = st.text_area(
    "Additional Context (optional)",
    placeholder="Any specific context for the LLM to consider...\ne.g. Company recently lost market share in rural areas. CEO is focused on EV transition. Meeting is with the CMO.",
    height=120
)

sources = st.text_area(
    "Reference Sources (optional)",
    placeholder="Sources to reference as starting points (LLM will expand on these)...\ne.g. McKinsey CX report 2023, Screener.in financials, SIAM annual report",
    height=80
)

# ---- SECTIONS SELECTOR ----
st.markdown('<div class="section-header">📋 Sections to Include</div>', unsafe_allow_html=True)

col5, col6, col7, col8, col9 = st.columns(5)
with col5:
    inc_industry = st.checkbox("Industry Overview & Trends", value=True)
with col6:
    inc_competitors = st.checkbox("Competitive Landscape", value=True)
with col7:
    inc_company = st.checkbox("Company Performance", value=True)
with col8:
    inc_priorities = st.checkbox("Strategic Priorities", value=True)
with col9:
    inc_ai = st.checkbox("AI Use Cases", value=True)

# ---- GENERATE BUTTON ----
st.markdown("---")
generate = st.button("🚀 Generate Research")

# ---- RESEARCH & DOC GENERATION ----
if generate:
    # Validations
    if not api_key:
        st.error("Please enter your Gemini API key in the sidebar.")
        st.stop()
    if not company or not industry or not meeting_topic:
        st.error("Please fill in Company Name, Industry and Meeting Topic.")
        st.stop()
    if not competitors_raw:
        st.error("Please enter at least one competitor.")
        st.stop()

    competitors = [c.strip() for c in competitors_raw.split(',') if c.strip()]

    # Geography string
    if geography_option == "India Only":
        geography = "India"
    elif geography_option == "India + Global":
        geography = "India and global markets"
    else:
        geography = f"India and {specific_countries}"

    # Configure Gemini
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')

    def ask_gemini(prompt):
        return model.generate_content(prompt).text

    CONTEXT = f"""
Company: {company}
Industry: {industry}
Competitors: {', '.join(competitors)}
Meeting Topic: {meeting_topic}
Meeting Objective: {meeting_objective}
Geography Focus: {geography}
Additional Context: {context_notes if context_notes else 'None provided'}
Reference Sources: {sources if sources else 'None provided – use best available sources'}

CRITICAL OUTPUT RULES:
- This is prep material for a senior consulting partner. Deck language only.
- No greetings, no "Dear CEO", no "In conclusion", no obvious statements
- Short punchy sentences. Numbers and specific examples wherever possible
- All insights directly relevant to the topic: "{meeting_topic}"
- Number all key points. Cluster related points under clear theme headers
- Always cite data source and year e.g. (Screener, FY24) or (McKinsey, 2023)
- Geography focus: {geography}
- Where reference sources are provided, use them as starting points and expand
"""

    sections_to_run = []
    if inc_industry:
        sections_to_run.append("industry")
    if inc_competitors:
        sections_to_run.append("competitors")
    if inc_company:
        sections_to_run.append("company")
    if inc_priorities:
        sections_to_run.append("priorities")
    if inc_ai:
        sections_to_run.append("ai")

    total = len(sections_to_run)
    progress = st.progress(0)
    status = st.empty()
    results = {}

    for i, section in enumerate(sections_to_run):
        status.info(f"Researching section {i+1} of {total}...")

        if section == "industry":
            results["industry"] = ask_gemini(f"""
{CONTEXT}
Research the {industry} industry focused on "{meeting_topic}" in {geography}.
Structure response under:
## Market Context
- Market size, growth rate, 2-3 key facts relevant to {meeting_topic}
## Key Trends (number 1-5)
For each: trend name, 2-3 bullet facts with numbers, one cross-industry insight
## Tailwinds for {meeting_topic}
- 3-4 numbered points with data
## Headwinds & Risks
- 3-4 numbered points with data
## Government & Regulatory Context
- 2-3 most relevant policy points
""")

        elif section == "competitors":
            results["competitors"] = ask_gemini(f"""
{CONTEXT}
Analyse how {', '.join(competitors)} approach "{meeting_topic}" vs {company} in {geography}.
Structure response under:
## Financial Benchmarking Table
Markdown table with columns:
| Company | Rev FY24 | 3Y Rev Growth | EBITDA FY24 | PAT FY24 | Mkt Share FY24 | Mkt Share Change |
Include all companies including {company}.
## Competitive Comparison: {meeting_topic}
Markdown table comparing all companies across 6 dimensions most relevant to {meeting_topic}.
| Dimension | {company} | {' | '.join(competitors)} |
## Key Competitive Insights (number 1-5)
Most important observations. Max 2 lines each. Specific examples with numbers.
""")

        elif section == "company":
            results["company"] = ask_gemini(f"""
{CONTEXT}
Analyse {company}'s performance focused on "{meeting_topic}" in {geography}.
Structure response under:
## Financial Performance Table
Markdown table:
| Metric | FY22 | FY23 | FY24 | Q3 FY24 | Q4 FY24 | YoY Q4 | QoQ |
Include: Revenue, EBITDA Margin, PAT, PAT Margin, Volumes, Avg Selling Price
## Performance vs Peers on {meeting_topic}
2-3 bullet points where {company} leads and lags
## Management Priorities on {meeting_topic}
Numbered points. Max 2 lines each. Include quotes where available.
## Key Challenges (number 1-4)
Direct and specific. Use data.
## Analyst View
Consensus rating, target price range, top 3 concerns relevant to {meeting_topic}
""")

        elif section == "priorities":
            results["priorities"] = ask_gemini(f"""
{CONTEXT}
Generate strategic recommendations for {company} on "{meeting_topic}" in {geography}.
Structure response under:
## Strategic Priority Areas
4-5 numbered priorities. For each:
### Priority N: [Name]
**Why critical:** 1-2 lines with data
**Gap vs best-in-class:** specific example from peer or another industry
**Recommended initiatives:**
  1. Initiative – what to do, expected impact, example company
  2. Initiative – same format
  3. Initiative – same format
**Risk of inaction:** 1 line
## Quick Wins (0-6 months)
3 numbered actions {company} can take immediately
## Medium Term (6-18 months)
3 numbered initiatives requiring more investment
""")

        elif section == "ai":
            results["ai"] = ask_gemini(f"""
{CONTEXT}
Generate AI use cases for {company} focused on "{meeting_topic}" in {geography}.
Structure response under:
## AI Use Case Themes
Exactly 4 themes relevant to {meeting_topic}. For each:
### Theme N: [Theme Name]
**Use Case 1:** Name
- What it does: 1 line
- Example: real company + result
- Impact: estimated % uplift or saving
- Complexity: Low / Medium / High
**Use Case 2:** same format
**Use Case 3:** same format
## Foundational Capabilities Required (number 1-3)
Specific – data infrastructure, talent, partnerships
## Recommended Starting Point
Single highest-impact, lowest-complexity use case and why
""")

        progress.progress((i + 1) / total)

    status.success("✅ Research complete! Generating document...")

    # ---- DOCUMENT GENERATION ----
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_divider(doc, hex_color='1F3564'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), hex_color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def parse_and_write(doc, raw_text):
        lines = raw_text.split('\n')
        current_table = None
        for line in lines:
            line_s = line.strip()
            if not line_s:
                if current_table is None:
                    doc.add_paragraph()
                continue
            if line_s.startswith('|'):
                if re.match(r'^\|[\s\-\|:]+\|$', line_s):
                    continue
                cells = [c.strip() for c in line_s.strip('|').split('|')]
                if current_table is None:
                    current_table = doc.add_table(rows=1, cols=len(cells))
                    current_table.style = 'Table Grid'
                    hdr = current_table.rows[0]
                    for i, ct in enumerate(cells):
                        hdr.cells[i].text = ct
                        set_cell_bg(hdr.cells[i], '1F3564')
                        for para in hdr.cells[i].paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.size = Pt(9)
                                run.font.name = 'Arial'
                else:
                    row = current_table.add_row()
                    ridx = len(current_table.rows) - 1
                    bg = 'EAF2FB' if ridx % 2 == 0 else 'FFFFFF'
                    for i, ct in enumerate(cells):
                        row.cells[i].text = ct
                        set_cell_bg(row.cells[i], bg)
                        for para in row.cells[i].paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(9)
                                run.font.name = 'Arial'
                continue
            else:
                current_table = None

            if line_s.startswith('## '):
                p = doc.add_heading(line_s[3:].strip(), level=2)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(13)
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                add_divider(doc, '2E74B5')
            elif line_s.startswith('### '):
                p = doc.add_heading(line_s[4:].strip(), level=3)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)
            elif line_s.startswith(('- ', '* ')):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(line_s[2:].replace('**','').strip())
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            elif re.match(r'^\d+\.', line_s):
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(line_s.replace('**','').strip())
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            elif line_s.startswith('**') and line_s.endswith('**'):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(line_s.replace('**','').strip())
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
            else:
                text = line_s.replace('**','').replace('*','').strip()
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run(text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(f"Strategic Overview: {company}")
    run.font.name = 'Arial'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Topic: {meeting_topic.title()}  |  Objective: {meeting_objective}")
    run2.font.name = 'Arial'
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0xC0, 0x50, 0x20)

    p3 = doc.add_paragraph()
    run3 = p3.add_run(f"Industry: {industry}  |  Geography: {geography}  |  March 2026  |  Confidential")
    run3.font.name = 'Arial'
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    add_divider(doc, '1F3564')

    section_map = [
        ("industry",    "1. Industry Overview & Trends"),
        ("competitors", "2. Competitive Landscape"),
        ("company",     "3. Company Performance & Challenges"),
        ("priorities",  "4. Strategic Priorities"),
        ("ai",          "5. AI Use Cases for Growth"),
    ]

    for key, title in section_map:
        if key in results:
            doc.add_page_break()
            p = doc.add_heading(title, level=1)
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
            add_divider(doc, '1F3564')
            doc.add_paragraph()
            parse_and_write(doc, results[key])

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    status.success("✅ Done! Your document is ready.")
    progress.progress(1.0)

    company_slug = company.replace(' ', '_')
    topic_slug = meeting_topic.replace(' ', '_')
    filename = f"{company_slug}_{topic_slug}_Research.docx"

    st.download_button(
        label="📥 Download Research Document",
        data=buf,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
